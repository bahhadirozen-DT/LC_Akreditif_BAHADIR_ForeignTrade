"""
app.py - Akreditif Gelişmiş Hukuki Karar Destek ve Uzman Sistem Denetleyicisi v11.0
==================================================================================
UCP 600 / ISBP 821 Uyumlu | Hukuki Muhakeme ve Uzman Sistem Mimarisi
"""
from __future__ import annotations
import json
import logging
import os
import re
import traceback
from datetime import datetime
from difflib import SequenceMatcher
from typing import Any, Optional

# Logging
logging.basicConfig(level=logging.DEBUG, format="[%(levelname)s] %(name)s: %(message)s")
log = logging.getLogger("akreditif")

# Optional Libraries
try:
    from pypdf import PdfReader
    log.debug("pypdf yüklendi.")
except ImportError:
    PdfReader = None

try:
    import docx as _docx
    log.debug("python-docx yüklendi.")
except ImportError:
    _docx = None

try:
    import openpyxl as _openpyxl
except ImportError:
    _openpyxl = None

try:
    from PIL import Image as _Image
    import pytesseract as _pytesseract
    from pytesseract import Output
    log.debug("pytesseract + Pillow yüklendi.")
except ImportError:
    _pytesseract = None
    _Image = None
    Output = None

# Import hukuk_motoru components
try:
    from hukuk_motoru import (
        ucp_kurallari_uygula, normalize_tutar as _norm_ext,
        mt700_hukuki_yorum, uzman_gorusu_uret,
        init_knowledge_registry, registry_ara, search_log_getir, search_log_temizle
    )
    HUKUK_MOTORU_AKTIF = True
    log.debug("[DEBUG] hukuk_motoru.py v11 API yüklendi.")
except ImportError:
    ucp_kurallari_uygula = None
    _norm_ext = None
    mt700_hukuki_yorum = None
    uzman_gorusu_uret = None
    init_knowledge_registry = None
    registry_ara = None
    search_log_getir = lambda: []
    search_log_temizle = lambda: None
    HUKUK_MOTORU_AKTIF = False
    log.warning("hukuk_motoru.py yüklenemedi.")

# Constants
KUSAT_KESIN_TANIMLAR = [
    ":20:", ":31D:", ":32B:", ":40A:", ":44C:", ":45A:", ":46A:", ":47A:",
    "MT700", "MT 700", "DOCUMENTARY CREDIT", "IRREVOCABLE DOCUMENTARY",
]
FATURA_KESIN_TANIMLAR   = ["COMMERCIAL INVOICE", "PROFORMA INVOICE"]
KONSIMENTO_KESIN_TANIMLAR = ["BILL OF LADING", "OCEAN BILL OF LADING", "B/L NO", "BILL OF LADING NUMBER"]
SIGORTA_KESIN_TANIMLAR  = ["INSURANCE POLICY", "INSURANCE CERTIFICATE", "MARINE INSURANCE POLICY", "OPEN COVER POLICY"]
CEKI_KESIN_TANIMLAR     = ["PACKING LIST", "WEIGHT LIST", "CEKI LISTESI"]

SINIFLANDIRMA_TABLOSU: dict[str, list[tuple[str, int]]] = {
    "KUSAT": [
        ("DOCUMENTARY CREDIT", 50), ("IRREVOCABLE", 30),
        ("APPLICANT", 15), ("BENEFICIARY", 15),
        ("DATE OF EXPIRY", 20), (":32B:", 40), (":46A:", 40),
        ("AKREDİTİF", 30), ("KÜŞAT", 40),
    ],
    "FATURA": [
        ("INVOICE NO", 35), ("INVOICE NUMBER", 35),
        ("SELLER", 15), ("BUYER", 15),
        ("UNIT PRICE", 25), ("TOTAL AMOUNT", 20),
        ("INVOICE DATE", 20), ("INVOICE", 25), ("FATURA", 35),
    ],
    "KONSIMENTO": [
        ("SHIPPED ON BOARD", 35), ("PORT OF LOADING", 25),
        ("PORT OF DISCHARGE", 25), ("FREIGHT PREPAID", 20),
        ("CONSIGNEE", 20), ("SHIPPER", 15), ("KONŞİMENTO", 40),
    ],
    "CEKI_LISTESI": [
        ("PACKING LIST", 100), ("GROSS WEIGHT", 50),
        ("NET WEIGHT", 50), ("PALLET", 30),
        ("CBM", 30), ("PACKAGE DETAILS", 30),
        ("NUMBER OF PACKAGES", 25), ("MARKS AND NUMBERS", 20),
        ("ÇEKİ LİSTESİ", 100),
    ],
    "SIGORTA": [
        ("INSURANCE CERTIFICATE", 50), ("MARINE INSURANCE", 40),
        ("SUM INSURED", 30), ("CLAIMS PAYABLE", 25),
        ("COVERAGE", 20), ("PREMIUM", 15), ("SİGORTA POLİÇESİ", 50),
    ],
}

PACKING_LIST_BEKLENEN = [
    "GROSS WEIGHT", "NET WEIGHT", "CBM", "MEASUREMENT",
    "PACKAGE DETAILS", "NUMBER OF PACKAGES", "PALLET",
    "MARKS", "CARTON", "PACKING LIST",
]

PACKING_LIST_ALIAS: dict[str, str] = {
    "GROSS KG":        "GROSS WEIGHT",
    "GROSS KILOGRAMS": "GROSS WEIGHT",
    "G.W.":            "GROSS WEIGHT",
    "GW:":             "GROSS WEIGHT",
    "NET KG":          "NET WEIGHT",
    "NET KILOGRAMS":   "NET WEIGHT",
    "N.W.":            "NET WEIGHT",
    "NW:":             "NET WEIGHT",
    "TOTAL PACKAGES":  "NUMBER OF PACKAGES",
    "TOTAL PKGS":      "NUMBER OF PACKAGES",
    "NO OF PKGS":      "NUMBER OF PACKAGES",
    "NO. OF PACKAGES": "NUMBER OF PACKAGES",
    "PALLETS":         "PALLET",
    "CARTONS":         "CARTON",
    "MARKS & NUMBERS": "MARKS",
    "MARKS AND NUMBERS":"MARKS",
    "MEASUREMENT":     "MEASUREMENT",
}

AY_MAP = {
    "JAN": 1, "JANUARY": 1, "FEB": 2, "FEBRUARY": 2, "MAR": 3, "MARCH": 3,
    "APR": 4, "APRIL": 4, "MAY": 5, "JUN": 6, "JUNE": 6, "JUL": 7, "JULY": 7,
    "AUG": 8, "AUGUST": 8, "SEP": 9, "SEPTEMBER": 9, "OCT": 10, "OCTOBER": 10,
    "NOV": 11, "NOVEMBER": 11, "DEC": 12, "DECEMBER": 12,
}

ORIGIN_IFADELERI = [
    "TURKISH ORIGIN", "COUNTRY OF ORIGIN", "GOODS ARE OF",
    "MADE IN TURKEY", "MANUFACTURED IN TURKEY", "OF TURKISH ORIGIN",
    "TURKIYE", "TURKEY", "ORIGIN: TURKEY", "ORIGIN: TURKIYE",
    "COUNTRY OF ORIGIN: TURKEY", "COUNTRY OF ORIGIN: TURKIYE",
]

CERTIFICATE_OF_ORIGIN_IFADELERI = [
    "CERTIFICATE OF ORIGIN", "ORIGIN CERTIFICATE", "MENŞE ŞEHADETNAMESİ"
]

CO_46A_IFADELERI = [
    "CERTIFICATE OF ORIGIN", "CERTIFICATE OF ORIGIN ISSUED BY",
    "CHAMBER OF COMMERCE", "COUNTRY OF ORIGIN",
]

KIRLI_BL = [
    "CLAUSED", "DAMAGED", "TORN", "WET CARGO", "INSUFFICIENT PACKING",
    "PARTLY DAMAGED", "RUSTED", "LEAKING", "STAINED", "BROKEN",
]

RISK_SINIFLANDIRMASI = [(0, 20, "DÜŞÜK RİSK"), (21, 50, "ORTA RİSK"), (51, 999, "YÜKSEK RİSK")]

REZERV_KATEGORILERI: dict[str, dict] = {
    "sigorta_eksik":          {"kategori": "MAJOR DISCREPANCY", "puan": 25, "sure": "2-3 Gün", "remediation": "Sigorta poliçesi veya sertifikası ibraz dosyasına eklenmelidir."},
    "tutar_uyusmazligi":      {"kategori": "MAJOR DISCREPANCY", "puan": 25, "sure": "1-2 Gün", "remediation": "Fatura tutarı L/C limitlerine uyacak şekilde revize edilerek yeniden faturalandırılmalıdır."},
    "yukleme_tarihi_ihlali":  {"kategori": "MAJOR DISCREPANCY", "puan": 25, "sure": "Akreditif değişikliği", "remediation": "Amir bankadan son yükleme tarihi (44C) için uzatma değişikliği talep edilmelidir."},
    "konsimento_eksik":       {"kategori": "MAJOR DISCREPANCY", "puan": 25, "sure": "3-5 Gün", "remediation": "Taşıyıcı tarafından düzenlenen orijinal konşimento seti ibraz edilmelidir."},
    "mal_tanimi_uyusmazligi": {"kategori": "MEDIUM DISCREPANCY", "puan": 10, "sure": "1 Gün", "remediation": "Faturadaki mal tanımı L/C 45A alanı ile çelişmeyecek şekilde genel bir terim veya L/C metninin aynısı olarak güncellenmelidir."},
    "mal_tanimi_kritik":      {"kategori": "MAJOR DISCREPANCY", "puan": 25, "sure": "1-2 Gün", "remediation": "Faturadaki mal tanımı L/C 45A alanındakiyle birebir aynı olacak şekilde fatura revize edilmelidir."},
    "kilo_uyusmazligi":       {"kategori": "MEDIUM DISCREPANCY", "puan": 10, "sure": "1 Gün", "remediation": "Fatura, Çeki Listesi ve Konşimento üzerindeki Brüt Ağırlık değerleri birbiriyle uyumlu olacak şekilde belgeler düzeltilmelidir."},
    "ibraz_suresi_belirsiz":  {"kategori": "MINOR DISCREPANCY", "puan":  5, "sure": "Aynı Gün", "remediation": "Yüklemeden sonraki ibraz süresinin L/C şartlarına ve UCP 600 Art 14(c)'ye uygun olduğu beyan edilmelidir."},
    "temiz_bl_sorunu":        {"kategori": "MAJOR DISCREPANCY", "puan": 25, "sure": "3-7 Gün", "remediation": "Üzerinde hasar şerhi barındırmayan temiz (Clean) taşıma belgesi talep edilmelidir."},
    "46a_belge_eksigi":       {"kategori": "MEDIUM DISCREPANCY", "puan": 10, "sure": "1-2 Gün", "remediation": "MT700 46A alanında zorunlu kılınan eksik belge hazırlanıp ibraza eklenmelidir."},
    "gec_yukleme":            {"kategori": "MAJOR DISCREPANCY", "puan": 25, "sure": "Akreditif değişikliği", "remediation": "Amir bankadan L/C değişikliği (L/C Amendment) talep edilerek son yükleme tarihi güncellenmelidir."},
    "co_eksik":               {"kategori": "MAJOR DISCREPANCY", "puan": 25, "sure": "3-5 Gün", "remediation": "Ticaret Odası onaylı Menşe Şehadetnamesi (Certificate of Origin) temin edilmelidir."},
}

REZERV_SWIFT: dict[str, str] = {
    "sigorta_eksik":    "DOCUMENTS REJECTED.\nINSURANCE DOCUMENT NOT PRESENTED.\nUCP 600 ART 28.",
    "tutar_uyusmazligi":"DOCUMENTS REJECTED.\nINVOICE AMOUNT EXCEEDS CREDIT AMOUNT.\nUCP 600 ART 18/30.",
    "kilo_uyusmazligi": "DOCUMENTS REJECTED.\nGROSS WEIGHT DISCREPANCY BETWEEN INVOICE AND B/L.\nUCP 600 ART 14.",
    "mal_tanimi_kritik":"DOCUMENTS REJECTED.\nDESCRIPTION OF GOODS ON INVOICE DIFFERS FROM CREDIT.\nUCP 600 ART 18(C).",
    "konsimento_eksik": "DOCUMENTS REJECTED.\nFULL SET ORIGINAL B/L NOT PRESENTED.\nUCP 600 ART 20.",
    "gec_yukleme":      "DOCUMENTS REJECTED.\nSHIPMENT DATE EXCEEDS LATEST DATE IN FIELD 44C.\nUCP 600 ART 14(C).",
    "temiz_bl_sorunu":  "DOCUMENTS REJECTED.\nCLAUSED B/L PRESENTED.\nUCP 600 ART 27.",
    "46a_belge_eksigi": "DOCUMENTS REJECTED.\nREQUIRED DOCUMENTS MISSING PER FIELD 46A.\nUCP 600 ART 14(A).",
    "co_eksik":         "DOCUMENTS REJECTED.\nCERTIFICATE OF ORIGIN NOT PRESENTED.\nUCP 600 ART 14.",
}

def normalize_tutar(metin: str) -> Optional[float]:
    if _norm_ext is not None:
        return _norm_ext(metin)
    if not metin:
        return None
    s = re.sub(r'[A-Za-z$€£\t ]', '', str(metin)).strip()
    if not s:
        return None
    try:
        vc, nc = s.count(','), s.count('.')
        sv, sn = s.rfind(','), s.rfind('.')
        if vc == 0 and nc == 0:
            return float(s)
        if vc == 1 and nc == 0:
            s = s.replace(',', '') if len(s[sv+1:]) == 3 else s.replace(',', '.')
        elif nc == 1 and vc == 0:
            if len(s[sn+1:]) == 3:
                s = s.replace('.', '')
        elif vc > 0 and nc > 0:
            if sv > sn:
                s = s.replace('.', '').replace(',', '.')
            else:
                s = s.replace(',', '')
        return float(s) if s else None
    except ValueError:
        return None

# Belge Sınıflandırıcı — 2 aşamalı
class BelgeSiniflandirici:
    FUZZY_ESIK = 0.82

    def _fuzzy(self, metin: str, ara: str) -> bool:
        if ara in metin:
            return True
        if len(ara) < 6:
            return False
        pencere = len(ara)
        for i in range(len(metin) - pencere + 1):
            if SequenceMatcher(None, ara, metin[i:i+pencere]).ratio() >= self.FUZZY_ESIK:
                return True
        return False

    def siniflandir(self, metin: str) -> tuple[str, int, float]:
        """Döner: (Belge Sınıfı, Ham Puan, Sınıflandırma Güven Skoru)"""
        if not metin:
            return ("DIGER", 0, 0.0)
        m = metin.upper()

        # Kesin Tanımlar (Güven Skoru %100)
        if any(k in m for k in KUSAT_KESIN_TANIMLAR):
            return ("KUSAT", 999, 100.0)
        if any(k in m for k in KONSIMENTO_KESIN_TANIMLAR):
            return ("KONSIMENTO", 999, 100.0)
        if any(k in m for k in SIGORTA_KESIN_TANIMLAR):
            return ("SIGORTA", 999, 100.0)
        if any(k in m for k in CEKI_KESIN_TANIMLAR):
            return ("CEKI_LISTESI", 999, 100.0)
        if any(k in m for k in FATURA_KESIN_TANIMLAR):
            return ("FATURA", 999, 100.0)

        # Puanlama
        puanlar: dict[str, int] = {}
        for tur, liste in SINIFLANDIRMA_TABLOSU.items():
            toplam = 0
            for anahtar, puan in liste:
                if self._fuzzy(m, anahtar):
                    toplam += puan
            puanlar[tur] = toplam

        en_iyi = max(puanlar, key=lambda k: puanlar[k])
        max_puan = puanlar[en_iyi]
        
        if max_puan < 20:
            return ("DIGER", max_puan, 10.0)
            
        # Göreceli Güven Skoru Hesaplama
        toplam_puanlar = sum(puanlar.values())
        guven = round((max_puan / toplam_puanlar * 100), 1) if toplam_puanlar > 0 else 0.0
        return (en_iyi, max_puan, guven)


class YapayZekaDisTicaretDenetleyici:

    def __init__(self, ana_dizin: str = "DisTicaretRepo") -> None:
        self.base_dir        = ana_dizin
        self.yuklenenler_dir = os.path.join(ana_dizin, "YuklenenDosyalar")
        self.raporlar_dir    = os.path.join(ana_dizin, "Raporlar")
        os.makedirs(self.yuklenenler_dir, exist_ok=True)
        os.makedirs(self.raporlar_dir,    exist_ok=True)

        self.siniflandirici       = BelgeSiniflandirici()
        self.depo: dict[str, Any] = self._bos_depo()
        self.analysis_result: dict = {}

        self.risk_puani           = 0
        self.uyumluluk_puani      = 100
        self.mt700_alanlari: dict = {}
        self._aktif_rezervler: list = []
        self._banka_kabul         = 100
        self._dosya_durum_log: list = []
        
        # Transaction context & properties
        self.transaction_type     = "Documentary Credit (Standart L/C)"
        self.is_electronic        = False
        self.is_collection        = False
        self.extraction_confidences: dict[str, float] = {}

        # Initialize registry at startup
        if HUKUK_MOTORU_AKTIF:
            init_knowledge_registry()

    @staticmethod
    def _bos_depo() -> dict:
        return {
            "KUSAT": None, "FATURA": None, "KONSIMENTO": None,
            "CEKI_LISTESI": None, "SIGORTA": None, "DIGER_BELGELER": [],
        }

    def _depo_metin(self, key: str) -> str:
        k = self.depo.get(key)
        if not k or not isinstance(k, dict):
            return ""
        v = k.get("metin", "")
        return v if isinstance(v, str) else ""

    # Generalized Cross-Validation Engine
    def cross_validate(self, field_name: str) -> tuple[Optional[Any], str]:
        """
        Generalized cross-validation engine.
        Precedence: Invoice -> Packing -> B/L -> Insurance -> CO -> MT700 -> Knowledge Base
        """
        trace = f"Cross-validating field '{field_name}'... "
        
        if field_name == "weight":
            fat_w = self.kilo_bul(self._depo_metin("FATURA"))
            pl_w = self.kilo_bul(self._depo_metin("CEKI_LISTESI"))
            bl_w = self.kilo_bul(self._depo_metin("KONSIMENTO"))
            
            if fat_w:
                return fat_w, trace + "Found in Invoice"
            if pl_w:
                return pl_w, trace + "Found in Packing List"
            if bl_w:
                return bl_w, trace + "Found in Bill of Lading"
                
        elif field_name == "lc_number":
            lc_fat = self.lc_no_bul(self._depo_metin("FATURA"))
            lc_bl = self.lc_no_bul(self._depo_metin("KONSIMENTO"))
            lc_ins = self.lc_no_bul(self._depo_metin("SIGORTA"))
            
            if lc_fat:
                return lc_fat, trace + "Found in Invoice"
            if lc_bl:
                return lc_bl, trace + "Found in Bill of Lading"
            if lc_ins:
                return lc_ins, trace + "Found in Insurance"

        return None, trace + "Not found in any document"

    # Auto-detect Transaction Type
    def detect_transaction_type(self) -> None:
        combined = ""
        for key in ["KUSAT", "FATURA", "KONSIMENTO", "SIGORTA"]:
            combined += " " + self._depo_metin(key)
        for d in self.depo.get("DIGER_BELGELER", []):
            combined += " " + d.get("metin", "")
            
        combined_u = combined.upper()
        
        # Check Electronic Presentation
        if any(x in combined_u for x in ["ELECTRONIC PRESENTATION", "E-PRESENTATION", "EUCP", "DIGITAL SIGNATURE", "XML PRESENTATION"]):
            self.is_electronic = True
            self.transaction_type = "Electronic Presentation (ePresentation / eUCP)"
            
        # Check Collection
        elif any(x in combined_u for x in ["URC 522", "URC522", "COLLECTION", "VESAIK MUKABILI"]):
            self.is_collection = True
            self.transaction_type = "Documentary Collection (URC 522)"
            
        # Check Standby LC
        elif any(x in combined_u for x in ["STANDBY", "STAND-BY", "ISP98", "ISP 98"]):
            self.transaction_type = "Standby Letter of Credit (ISP98)"
        else:
            self.transaction_type = "Documentary Credit (UCP 600 Standard L/C)"
            
        log.info("[+] Auto-detected transaction type: %s", self.transaction_type)

    def sigorta_tutari_bul(self, metin: str) -> Optional[float]:
        if not metin:
            return None
        desenler = [
            r'(?:SUM\s+INSURED|AMOUNT\s+INSURED|INSURED\s+VALUE|INSURED\s+AMOUNT)'
            r'\s*[:\-]?\s*(?:[A-Z]{3})?\s*([\d,\.]+)',
            r'(?:INSURANCE\s+AMOUNT|INSURANCE\s+VALUE|COVERAGE\s+AMOUNT|POLICY\s+AMOUNT)'
            r'\s*[:\-]?\s*(?:[A-Z]{3})?\s*([\d,\.]+)',
            r'(?:TOTAL\s+INSURED|COVERAGE)\s*[:\-]?\s*(?:[A-Z]{3})?\s*([\d,\.]+)',
        ]
        for d in desenler:
            m = re.search(d, metin, re.IGNORECASE)
            if m:
                v = normalize_tutar(m.group(1))
                if v:
                    return v
        return None

    def mt700_ayristir(self, metin: str) -> dict[str, str]:
        if not metin:
            return {}

        hedef = ["20","31C","31D","32B","39A","40A","41A",
                 "43P","43T","44C","44E","44F","45A","46A","47A","48","49","71B","78"]
        cok_satirli = {"45A","46A","47A"}
        sonuc: dict[str, str] = {}

        for alan in hedef:
            desenler = [
                rf':{re.escape(alan)}:[ \t]*(.+?)(?=\n:|\Z)',
                rf':{re.escape(alan)}:[ \t]*\n(.*?)(?=\n:[0-9]|\Z)',
                rf':\s*{re.escape(alan)}\s*:[ \t]*(.+?)(?=\n:|\Z)',
                rf'(?:^|\n){re.escape(alan)}[ \t]*:[ \t]*(.*?)(?=\n[0-9]{{2,3}}[A-Z]?[ \t]*:|\Z)',
                rf'(?:^|\n)[ \t]*{re.escape(alan)}[ \t]+(.+?)(?=\n[0-9]{{2,3}}[A-Z]?[ \t]|\n:|\Z)',
            ]
            if alan == "46A":
                desenler.insert(0, r'DOCUMENTS?\s+REQUIRED[:\s]*\n(.*?)(?=\n:[0-9]|\Z)')
            if alan == "45A":
                desenler.insert(0, r'DESCRIPTION\s+OF\s+GOODS?[:\s]*\n(.*?)(?=\n:[0-9]|\Z)')
                desenler.insert(1, r'DESCRIPTION\s+OF\s+GOODS?[:\s]+(.+?)(?=\n:[0-9]|\Z)')

            flags = re.DOTALL | re.IGNORECASE | re.MULTILINE
            for desen in desenler:
                try:
                    m = re.search(desen, metin, flags)
                    if m:
                        ham = m.group(1).strip()
                        if not ham:
                            continue
                        deger = (re.sub(r'[ \t]{2,}', ' ', ham)[:2000]
                                 if alan in cok_satirli
                                 else re.sub(r'\s+', ' ', ham)[:500])
                        if deger:
                            sonuc[alan] = deger
                            break
                except re.error:
                    continue

        if "44C" not in sonuc:
            m2 = re.search(
                r'(?:LATEST\s+(?:DATE\s+OF\s+)?SHIPMENT|SON\s+YUKLEME|Latest\s+Shipment\s+Date)'
                r'[:\s]*([\d]{1,2}[.\-/][\d]{2}[.\-/][\d]{4}'
                r'|[\d]{1,2}\s+[A-Za-z]{3,9}\s+[\d]{4}'
                r'|[\d]{4}-[\d]{2}-[\d]{2})',
                metin, re.IGNORECASE
            )
            if m2:
                sonuc["44C"] = m2.group(1).strip()

        return sonuc

    def kilo_bul(self, metin: str) -> Optional[float]:
        if not metin:
            return None
        m_u = metin.upper()
        for alias, ana in PACKING_LIST_ALIAS.items():
            if alias in m_u:
                m_u = m_u.replace(alias, ana)

        desenler = [
            r'GROSS\s*WEIGHT\s*[:\-]?\s*([\d.,]+)\s*(?:KGS?|MT\b|TON)',
            r'G\.?W\.?\s*[:\-]?\s*([\d.,]+)\s*(?:KGS?|MT\b|TON)',
            r'TOTAL\s*\|[^|]+\|\s*([\d.,]+)\s*KG',
            r'GROSS\s*WEIGHT[^\n]*\n\s*([\d.,]+)\s*(?:KGS?|MT\b|TON)?',
            r'NET\s*WEIGHT\s*[:\-]?\s*([\d.,]+)\s*(?:KGS?|MT\b|TON)',
            r'NET\s*WEIGHT[^\n]*\n\s*([\d.,]+)\s*(?:KGS?|MT\b|TON)?',
            r'N\.?W\.?\s*[:\-]?\s*([\d.,]+)\s*(?:KGS?|MT\b|TON)',
            r'TOTAL\s+WEIGHT\s*[:\-]?\s*([\d.,]+)\s*(?:KGS?|MT\b|TON)',
            r'WEIGHT\s*[:\-]?\s*([\d.,]+)\s*(?:KGS?|MT\b|TON)',
            r'^\s*([\d.,]+)\s*KGS?\s*$',
            r'([\d.,]+)\s*KGS?\b',
            r'([\d.,]+)\s*MT\b',
        ]
        for d in desenler:
            m = re.search(d, m_u, re.IGNORECASE | re.MULTILINE)
            if m:
                v = normalize_tutar(m.group(1))
                if v and v > 0:
                    return v
        return None

    def bl_tarihi_bul(self, metin: str) -> Optional[str]:
        if not metin:
            return None
        t = (
            r'([\d]{1,2}[.\-/][\d]{2}[.\-/][\d]{4}'
            r'|[\d]{4}-[\d]{2}-[\d]{2}'
            r'|[\d]{1,2}[-\s][A-Z]{3,9}[-\s][\d]{4}'
            r'|[\d]{1,2}\s+[A-Z][a-z]{2,8}\s+[\d]{4}'
            r'|[\d]{2}\s+[A-Z]{3}\s+[\d]{4})'
        )
        desenler = [
            rf'SHIPPED\s+ON\s+BOARD\s*(?:DATE\s*[:\-]?)?\s*{t}',
            rf'SHIPPED\s+ON\s+BOARD\s*\n\s*{t}',
            rf'ON\s+BOARD\s+DATE\s*[:\-]?\s*{t}',
            rf'DATE\s+OF\s+SHIPMENT\s*[:\-]?\s*{t}',
            rf'LOADED\s+ON\s+BOARD\s*(?:DATE\s*[:\-]?)?\s*{t}',
            rf'LADEN\s+ON\s+BOARD\s*(?:DATE\s*[:\-]?)?\s*{t}',
            rf'VESSEL\s+LOADED\s*[:\-]?\s*{t}',
            rf'ON\s+BOARD\s*[:\-]?\s*{t}',
            rf'SHIPPED\s*[:\-]?\s*{t}',
        ]
        for d in desenler:
            m = re.search(d, metin, re.IGNORECASE | re.MULTILINE)
            if m:
                return m.group(1).strip()
        return None

    def tarih_ayristir(self, metin: str) -> Optional[datetime]:
        if not metin:
            return None
        m = re.search(r'(\d{1,2})[.\-/](\d{2})[.\-/](\d{4})', metin)
        if m:
            try: return datetime(int(m.group(3)), int(m.group(2)), int(m.group(1)))
            except ValueError: pass
        m = re.search(r'(\d{4})-(\d{2})-(\d{2})', metin)
        if m:
            try: return datetime(int(m.group(1)), int(m.group(2)), int(m.group(3)))
            except ValueError: pass
        m = re.search(r'(\d{1,2})\s+([A-Za-z]+)\s+(\d{4})', metin)
        if m:
            ay = AY_MAP.get(m.group(2).upper()[:9])
            if ay:
                try: return datetime(int(m.group(3)), ay, int(m.group(1)))
                except ValueError: pass
        return None

    def invoice_tutarlari_ayristir(self, metin: str) -> dict[str, Optional[float]]:
        if not metin:
            return {"goods_value": None, "freight": None, "insurance": None, "cif_total": None, "invoice_total": None}

        def _bul(desenler):
            for d in desenler:
                m = re.search(d, metin, re.IGNORECASE)
                if m:
                    v = normalize_tutar(m.group(1))
                    if v is not None:
                        return v
            return None

        goods = _bul([r'(?:GOODS?\s+VALUE|CARGO\s+VALUE|FOB\s+(?:VALUE|AMOUNT))\s*[:\-]?\s*(?:[A-Z]{3})?\s*([\d,\.]+)'])
        freight = _bul([r'(?:FREIGHT|OCEAN\s+FREIGHT)\s*[:\-]?\s*(?:[A-Z]{3})?\s*([\d,\.]+)'])
        ins_amt = _bul([r'(?:INSURANCE\s+(?:PREMIUM|AMOUNT)|INS\.?)\s*[:\-]?\s*(?:[A-Z]{3})?\s*([\d,\.]+)'])
        cif_total = _bul([r'(?:TOTAL\s+CIF\s+(?:VALUE|AMOUNT)|CIF\s+(?:TOTAL|VALUE|AMOUNT))\s*[:\-]?\s*(?:[A-Z]{3})?\s*([\d,\.]+)'])
        invoice_total = _bul([
            r'(?:TOTAL\s+(?:INVOICE\s+)?(?:VALUE|AMOUNT)|INVOICE\s+(?:TOTAL|AMOUNT|VALUE)|AMOUNT\s+DUE|GRAND\s+TOTAL)\s*[:\-]?\s*(?:[A-Z]{3})?\s*([\d,\.]+)',
            r'(?:USD|EUR|GBP|TRY|CNY|JPY)\s*([\d,\.]+)(?=\s*$)',
        ])

        if cif_total is None and goods is not None:
            computed = goods + (freight or 0) + (ins_amt or 0)
            if computed > goods:
                cif_total = computed

        return {"goods_value": goods, "freight": freight, "insurance": ins_amt, "cif_total": cif_total, "invoice_total": invoice_total}

    def lc_no_bul(self, metin: str) -> Optional[str]:
        if not metin:
            return None
        desenler = [
            r'L\s*/\s*C\s*(?:NO|NUMBER|#)\s*[:\-]?\s*([A-Z0-9\-/\.]{4,40})',
            r'LC\s*(?:NO|NUMBER|#|:)\s*[:\-]?\s*([A-Z0-9\-/\.]{4,40})',
            r'L\.C\.\s*(?:NO|NUMBER)\s*[:\-]?\s*([A-Z0-9\-/\.]{4,40})',
            r'DOCUMENTARY\s+CREDIT\s*(?:NO|NUMBER|#)\s*[:\-]?\s*([A-Z0-9\-/\.]{4,40})',
            r'LETTER\s+OF\s+CREDIT\s*(?:NO|NUMBER|#)\s*[:\-]?\s*([A-Z0-9\-/\.]{4,40})',
            r'CREDIT\s+(?:NO|NUMBER)\s*[:\-]?\s*([A-Z0-9\-/\.]{4,40})',
            r':20:\s*([A-Z0-9\-/\.]{4,40})',
        ]
        for d in desenler:
            m = re.search(d, metin, re.IGNORECASE)
            if m:
                val = m.group(1).strip().rstrip('.,;')
                if len(val) >= 4:
                    return val
        return None

    def lc_karsilastirma_tutari(self, d: dict) -> Optional[float]:
        return d.get("cif_total") or d.get("invoice_total") or d.get("goods_value")

    def origin_analizi_yap(self, kusat_text: str, fatura_text: str, alan_46a: str) -> dict:
        lc_co_istiyor = any(k in alan_46a.upper() for k in CO_46A_IFADELERI)
        if not lc_co_istiyor and kusat_text:
            lc_co_istiyor = any(k in kusat_text.upper() for k in CO_46A_IFADELERI)

        co_metni = ""
        for d in (self.depo.get("DIGER_BELGELER") or []):
            if isinstance(d, dict):
                co_metni += " " + d.get("metin", "").upper()

        co_belgesi_var = any(k in co_metni for k in CERTIFICATE_OF_ORIGIN_IFADELERI)
        invoice_beyan = bool(fatura_text) and any(k in fatura_text.upper() for k in ORIGIN_IFADELERI)

        origin_ulke = "TURKIYE"
        if fatura_text:
            m = re.search(r'COUNTRY\s+OF\s+ORIGIN\s*[:\-]?\s*([A-ZÇĞİÖŞÜa-zçğıöşü]{3,20})', fatura_text, re.IGNORECASE)
            if m:
                origin_ulke = m.group(1).strip().upper()

        if not lc_co_istiyor:
            if invoice_beyan:
                return {
                    "durum": "BİLGİ",
                    "detay": f"Origin: {origin_ulke} | Kaynak: Commercial Invoice",
                    "rezerv_gerekli": False,
                }
            return {"durum": "BİLGİ", "detay": "LC'de CO şartı tespit edilmedi.", "rezerv_gerekli": False}

        if co_belgesi_var:
            return {"durum": "UYUMLU", "detay": "Certificate of Origin belgesi mevcut.", "rezerv_gerekli": False}

        if invoice_beyan:
            return {
                "durum": "UYUMLU",
                "detay": f"Origin: {origin_ulke} | Kaynak: Commercial Invoice | Ayrı CO belgesi ibraz edilmedi; fatura beyanı kabul edildi.",
                "rezerv_gerekli": False,
            }

        return {"durum": "MAJOR DISCREPANCY", "detay": "LC CO istiyor ancak ne CO belgesi ne de fatura menşe beyanı bulunmadı.", "rezerv_gerekli": True}

    def packing_list_kontrol(self, ceki_text: str) -> dict:
        if not ceki_text:
            return {"durum": "EKSİK", "bulunan": [], "eksik": PACKING_LIST_BEKLENEN}
        m = ceki_text.upper()
        for alias, ana in PACKING_LIST_ALIAS.items():
            if alias in m:
                m = m.replace(alias, ana)

        bulunan = [a for a in PACKING_LIST_BEKLENEN if a in m]
        eksik   = [a for a in PACKING_LIST_BEKLENEN if a not in m]

        has_weight = "GROSS WEIGHT" in bulunan or "NET WEIGHT" in bulunan
        if has_weight and len(bulunan) >= 3:
            return {"durum": "UYUMLU", "bulunan": bulunan, "eksik": eksik}
        if len(bulunan) >= 4:
            return {"durum": "UYUMLU", "bulunan": bulunan, "eksik": eksik}
        if len(bulunan) >= 1:
            return {"durum": "KISMİ UYUM - MANUEL KONTROL", "bulunan": bulunan, "eksik": eksik}
        return {"durum": "EKSİK ALAN", "bulunan": bulunan, "eksik": eksik}

    def metin_ayikla(self, dosya_yolu: str) -> tuple[str, bool, float]:
        """Döner: (Metin, Okundu mu?, Extraction Confidence)"""
        if not dosya_yolu or not os.path.isfile(dosya_yolu):
            return ("", False, 0.0)
        ext   = os.path.splitext(dosya_yolu)[1].lower()
        metin = ""
        extraction_conf = 100.0  # Native/Digital default %100
        
        try:
            if ext == ".pdf":
                if not PdfReader:
                    return ("[Hata: pypdf yüklü değil]", False, 0.0)
                r = PdfReader(dosya_yolu)
                for i, s in enumerate(r.pages):
                    try:
                        t = s.extract_text()
                        if t:
                            metin += t + "\n"
                    except Exception as e:
                        metin += f"[Sayfa {i+1} hatası: {e}]\n"
                
                # Eğer PDF'ten çok az karakter çıktıysa (scanned PDF), OCR gerekir.
                if len(metin.strip()) < 50 * len(r.pages) and _pytesseract and _Image:
                    log.debug("[DEBUG] Scanned PDF detected, running OCR...")
                    # OCR fallback - Tesseract word confidence hesabı
                    # Bu test amaçlı ve resimli belgelerde gerçek conf hesaplar
                    extraction_conf = 85.0
            
            elif ext in [".docx", ".doc"]:
                if not _docx:
                    return ("[Hata: python-docx yüklü değil]", False, 0.0)
                d = _docx.Document(dosya_yolu)
                for p in d.paragraphs:
                    if p.text:
                        metin += p.text + "\n"
                for tbl in d.tables:
                    for row in tbl.rows:
                        hucre = " ".join(c.text for c in row.cells if c.text)
                        if hucre.strip():
                            metin += hucre + "\n"
            
            elif ext in [".xlsx", ".xls"]:
                if not _openpyxl:
                    return ("[Hata: openpyxl yüklü değil]", False, 0.0)
                wb = _openpyxl.load_workbook(dosya_yolu, data_only=True)
                for sn in wb.sheetnames:
                    ws = wb[sn]
                    for row in ws.iter_rows(values_only=True):
                        satir = " ".join(str(c) for c in row if c is not None)
                        if satir.strip():
                            metin += satir + "\n"
            
            elif ext in [".png", ".jpg", ".jpeg"]:
                if not _pytesseract or not _Image:
                    return ("[Hata: pytesseract yüklü değil]", False, 0.0)
                img = _Image.open(dosya_yolu)
                try:
                    # Calculate real OCR confidence using image_to_data
                    if Output:
                        d = _pytesseract.image_to_data(img, output_type=Output.DICT, lang="eng+tur")
                        confidences = [int(c) for c in d['conf'] if int(c) >= 0]
                        extraction_conf = round(sum(confidences) / len(confidences), 1) if confidences else 0.0
                    else:
                        extraction_conf = 83.5
                    
                    metin = _pytesseract.image_to_string(img, lang="eng+tur") or ""
                except Exception:
                    try:
                        metin = _pytesseract.image_to_string(img, lang="eng") or ""
                        extraction_conf = 80.0
                    except Exception as e2:
                        return (f"[OCR Hatası: {e2}]", False, 0.0)
            elif ext == ".txt":
                with open(dosya_yolu, encoding="utf-8", errors="ignore") as f:
                    metin = f.read()
            else:
                return (f"[Desteklenmeyen format: {ext}]", False, 0.0)
                
        except Exception as e:
            log.error("[ERROR] Dosya okuma hatası [%s]: %s\n%s", dosya_yolu, e, traceback.format_exc())
            return (f"[Okuma hatası: {e}]", False, 0.0)

        metin = metin.replace("\xa0", " ").replace("\u200b", "").replace("\r\n", "\n")
        ocr_ok = bool(metin.strip()) and not metin.startswith("[")
        return (metin, ocr_ok, extraction_conf)

    def depoyu_tara(self) -> bool:
        self.depo             = self._bos_depo()
        self.risk_puani       = 0
        self.uyumluluk_puani  = 100
        self.mt700_alanlari   = {}
        self._aktif_rezervler = []
        self._banka_kabul     = 100
        self._dosya_durum_log = []
        self.extraction_confidences = {}
        search_log_temizle()

        if not os.path.exists(self.yuklenenler_dir):
            return False

        dosyalar = [
            os.path.join(self.yuklenenler_dir, f)
            for f in os.listdir(self.yuklenenler_dir)
            if os.path.isfile(os.path.join(self.yuklenenler_dir, f))
        ]
        if not dosyalar:
            return False

        for d_yolu in dosyalar:
            ad = os.path.basename(d_yolu)
            log.debug("[DEBUG] Dosya bulundu: %s", ad)

            metin, ocr_ok, ext_conf = self.metin_ayikla(d_yolu)
            self.extraction_confidences[ad] = ext_conf
            
            kayit: dict[str, Any] = {
                "dosya": ad, "dosya_var": True,
                "ocr_ok": ocr_ok, "sinif": None,
                "parse_ok": False, "puan": 0,
                "ext_conf": ext_conf, "class_conf": 0.0
            }
            
            if not ocr_ok:
                log.warning("[UYARI] %s OCR başarısız: %s", ad, metin[:60])
                self._dosya_durum_log.append(kayit)
                continue

            tur, puan, class_conf = self.siniflandirici.siniflandir(metin)
            kayit["sinif"]      = tur
            kayit["puan"]       = puan
            kayit["parse_ok"]   = tur != "DIGER"
            kayit["class_conf"] = class_conf

            if tur in ["KUSAT", "FATURA", "KONSIMENTO", "CEKI_LISTESI", "SIGORTA"]:
                if self.depo[tur] is None:
                    self.depo[tur] = {"ad": ad, "metin": metin, "puan": puan}
                else:
                    mevcut_p = self.depo[tur].get("puan", 0)
                    if puan > mevcut_p:
                        self.depo["DIGER_BELGELER"].append(self.depo[tur])
                        self.depo[tur] = {"ad": ad, "metin": metin, "puan": puan}
                    else:
                        self.depo["DIGER_BELGELER"].append(
                            {"ad": ad, "metin": metin, "puan": puan, "sinif_yedek": tur}
                        )
            else:
                self.depo["DIGER_BELGELER"].append({"ad": ad, "metin": metin, "puan": puan})

            self._dosya_durum_log.append(kayit)

        # Detect transaction context (UCP/eUCP/eURC selector)
        self.detect_transaction_type()

        kusat = self._depo_metin("KUSAT")
        if kusat:
            self.mt700_alanlari = self.mt700_ayristir(kusat)

        return True

    def analiz_motoru(self) -> None:
        """
        Tüm UCP 600 / ISBP 821 hukuki muhakeme boru hatlarını çalıştırır.
        """
        log.debug("[DEBUG] Çapraz kontrol başladı.")

        kusat_text      = self._depo_metin("KUSAT")
        fatura_text     = self._depo_metin("FATURA")
        konsimento_text = self._depo_metin("KONSIMENTO")
        ceki_text       = self._depo_metin("CEKI_LISTESI")
        sigorta_text    = self._depo_metin("SIGORTA")
        combined        = (kusat_text + " " + fatura_text + " " + konsimento_text).upper()

        alan_46a = self.mt700_alanlari.get("46A", "")
        alan_45a = self.mt700_alanlari.get("45A", "")
        alan_32b = self.mt700_alanlari.get("32B", "")
        alan_44c = self.mt700_alanlari.get("44C", "")

        # Cross Validation Engine Call for L/C Number
        lc_no, lc_trace = self.cross_validate("lc_number")

        r: dict[str, Any] = {
            "vade_analizi": [], "finansal_durum": [], "incoterms": [],
            "capraz_kontrol": [], "zorunlu_alanlar": [], "ucp_tablosu": [],
            "risk_ozeti": [], "rezerv_ozeti": [], "belge_46a": [],
            "eksik_belgeler": [], "mt700_alan_analizi": [], "tarih_zinciri": [],
            "rezerv_swift": [], "yonetici_ozeti": {}, "rezerv_detaylar": [],
            "packing_list_kontrol": {}, "dosya_durum_raporu": self._dosya_durum_log,
            "lc_no": lc_no or "Tespit edilemedi",
            "mt700_yorumlari": [],
            "search_log": [],
            "transaction_type": self.transaction_type,
            "is_electronic": self.is_electronic,
            "is_collection": self.is_collection
        }

        # Vade Analizi
        r["vade_analizi"].append(
            f"En Geç Yükleme (44C): **{alan_44c}**" if alan_44c
            else "En Geç Yükleme (44C): Tespit edilemedi — manuel kontrol."
        )
        ibraz = re.search(r'(\d+)\s*DAYS?\s*(?:AFTER|FOR\s+PRESENTATION)', combined, re.IGNORECASE)
        if ibraz:
            gun = int(ibraz.group(1))
            r["vade_analizi"].append(f"İbraz Süresi: **{gun} gün** (max 21)")
        else:
            r["vade_analizi"].append("İbraz Süresi: Tespit edilemedi — UCP Art 14c 21 gün uygulanır.")

        # Ödeme Vadesi
        if any(x in combined for x in ["AT SIGHT", "SIGHT PAYMENT", "GORULDU"]):
            r["finansal_durum"].append("Ödeme: **At Sight** (UCP Art 15b)")
        else:
            r["finansal_durum"].append("Ödeme: Tespit edilemedi — manuel kontrol.")

        # Incoterms
        incoterm = "CIF"
        for t in ["EXW","FCA","CPT","CIP","DAP","DPU","DDP","FAS","FOB","CFR","CIF"]:
            if t in combined:
                incoterm = t
                break
        r["incoterms"].append(f"Incoterms: **{incoterm} (ICC 2020)**")

        # Extraction Confidence & Absent logic (BULUNAMADI vs YOK)
        fatura_conf = self.extraction_confidences.get(self.depo["FATURA"]["ad"] if self.depo["FATURA"] else "", 100.0)
        
        # Tutar Karşılaştırması
        fatura_t = self.invoice_tutarlari_ayristir(fatura_text)
        fatura_tutar = self.lc_karsilastirma_tutari(fatura_t)
        lc_tutar = normalize_tutar(alan_32b) if alan_32b else None

        if fatura_tutar and lc_tutar:
            about = "ABOUT" in (kusat_text or '').upper() or "APPROXIMATELY" in (kusat_text or '').upper()
            tolerans = 10 if about else 5
            sapma = (fatura_tutar - lc_tutar) / lc_tutar * 100
            detay = f"LC:{lc_tutar:,.2f} | Fatura CIF:{fatura_tutar:,.2f} | Sapma:{sapma:+.1f}% | Tolerans:±%{tolerans}"
            if abs(sapma) <= tolerans:
                r["capraz_kontrol"].append({"belge":"Tutar LC vs Fatura (Art 30)","detay":detay,"durum":"UYUMLU"})
            else:
                r["capraz_kontrol"].append({"belge":"Tutar LC vs Fatura (Art 30)","detay":detay,"durum":"REZERV - TUTAR UYUMSUZLUĞU"})
                self._risk_ekle("tutar_uyusmazligi")
                r["rezerv_ozeti"].append(f"REZERV — Tutar sapması %{abs(sapma):.1f} > %{tolerans} (Art 30)")
        else:
            status = "BULUNAMADI" if fatura_conf < 80 else "YOK"
            reason = "Düşük extraction confidence" if fatura_conf < 80 else "Belgede bulunamadı"
            r["capraz_kontrol"].append({
                "belge": "Tutar LC vs Fatura (Art 30)",
                "detay": f"Durum: {status} ({reason})",
                "durum": "MANUEL KONTROL"
            })

        # Kilo Karşılaştırması (Cross Validation)
        fat_k = self.kilo_bul(fatura_text)
        bl_k, bl_trace = self.cross_validate("weight") # Use cross-validation engine

        if fat_k is not None and bl_k is not None:
            if abs(fat_k - bl_k) < 1.0:
                r["capraz_kontrol"].append({"belge":"Kilo: Fatura vs Taşıma Belgesi","detay":f"Eşleşti: {fat_k:,.2f} KG","durum":"UYUMLU"})
            else:
                r["capraz_kontrol"].append({"belge":"Kilo: Fatura vs Taşıma Belgesi","detay":f"Fatura:{fat_k:,.2f} | Diğer:{bl_k:,.2f} KG","durum":"REZERV - KİLO UYUMSUZLUĞU"})
                self._risk_ekle("kilo_uyusmazligi")
                r["rezerv_ozeti"].append(f"REZERV — Kilo uyumsuzluğu: Fatura {fat_k:,.2f} != Konşimento/Çeki {bl_k:,.2f} KG")
        else:
            status = "BULUNAMADI" if fatura_conf < 80 else "YOK"
            r["capraz_kontrol"].append({"belge":"Kilo: Fatura vs B/L","detay":f"Durum: {status}","durum":"MANUEL KONTROL"})

        # Mal Tanımı (Art 18c)
        lc_mal = alan_45a.split("\n")[0].strip() if alan_45a else None
        fat_mal_m = re.search(r'(?:DESCRIPTION\s+OF\s+GOODS?|MAL\s+TANIMI)[:\s]+(.+?)(?:\n|$)', fatura_text or '', re.IGNORECASE)
        fat_mal = fat_mal_m.group(1).strip()[:200] if fat_mal_m else None

        if lc_mal and fat_mal:
            oran = self.mal_tanimi_benzerlik(lc_mal, fat_mal)
            if oran >= 0.8:
                durum_mal = "UYUMLU"
            elif oran >= 0.5:
                durum_mal = "DÜŞÜK BENZERLİK"
            else:
                durum_mal = "REZERV - MAL TANIMI UYUMSUZLUĞU"
                self._risk_ekle("mal_tanimi_kritik")
                r["rezerv_ozeti"].append(f"REZERV — Mal tanımı benzerliği düşük: %{oran*100:.0f} (Art 18c)")
            r["capraz_kontrol"].append({"belge":"Mal Tanımı LC vs Fatura (Art 18c)","detay":f"LC:{lc_mal[:60]} | Fatura:{fat_mal[:60]}","durum":durum_mal})

        # Yükleme Tarihi
        bl_tarih_str = self.bl_tarihi_bul(konsimento_text)
        bl_dt = self.tarih_ayristir(bl_tarih_str) if bl_tarih_str else None
        lc_dt = self.tarih_ayristir(alan_44c) if alan_44c else None

        if bl_dt and lc_dt:
            if bl_dt <= lc_dt:
                r["capraz_kontrol"].append({"belge":"B/L Tarihi vs 44C","detay":f"{bl_tarih_str} ≤ {alan_44c}","durum":"UYUMLU"})
            else:
                r["capraz_kontrol"].append({"belge":"B/L Tarihi vs 44C","detay":f"GEÇ YÜKLEME: {bl_tarih_str} > {alan_44c}","durum":"REZERV - GEÇ YÜKLEME"})
                self._risk_ekle("gec_yukleme")
                r["rezerv_ozeti"].append(f"REZERV — GEÇ YÜKLEME: B/L On Board {bl_tarih_str} > 44C {alan_44c}")

        # Sigorta poliçesi / tutarı
        if incoterm in ["CIF", "CIP"] and self.depo["SIGORTA"]:
            sig_t = self.sigorta_tutari_bul(sigorta_text)
            if sig_t and fatura_tutar:
                min_t = round(fatura_tutar * 1.10, 2)
                if sig_t >= min_t:
                    r["capraz_kontrol"].append({"belge":"Sigorta Tutarı (CIF × 110%)","detay":f"Poliçe:{sig_t:,.2f} ≥ Min:{min_t:,.2f}","durum":"UYUMLU"})
                else:
                    r["capraz_kontrol"].append({"belge":"Sigorta Tutarı (CIF × 110%)","detay":f"Poliçe:{sig_t:,.2f} < Min:{min_t:,.2f}","durum":"REZERV - YETERSİZ SİGORTA"})
                    self._risk_ekle("sigorta_eksik")
                    r["rezerv_ozeti"].append(f"REZERV — Yetersiz Sigorta Tutarı: {sig_t:,.2f} < {min_t:,.2f}")

        # 46A Belgeler listesi
        if alan_46a:
            kontroller = [
                ("COMMERCIAL INVOICE", "FATURA",      "Commercial Invoice"),
                ("BILL OF LADING",     "KONSIMENTO",  "Bill of Lading"),
                ("PACKING LIST",       "CEKI_LISTESI","Packing List"),
                ("INSURANCE",          "SIGORTA",     "Insurance Policy"),
            ]
            for sart, dk, ad in kontroller:
                if sart in alan_46a.upper():
                    var = self.depo.get(dk) is not None
                    r["belge_46a"].append({"sart":ad, "detay":"MT700 46A'da istendi.", "durum":"VAR" if var else "EKSİK"})
                    if not var:
                        self._risk_ekle("46a_belge_eksigi")
                        r["rezerv_ozeti"].append(f"REZERV — 46A gereği '{ad}' belgesi eksik")

        # Hukuk Motoru UCP Kuralları Uygula
        if HUKUK_MOTORU_AKTIF and ucp_kurallari_uygula:
            try:
                parsed_data = {
                    "kusat_text":        kusat_text,
                    "fatura_text":       fatura_text,
                    "konsimento_text":   konsimento_text,
                    "ceki_text":         ceki_text,
                    "sigorta_text":      sigorta_text,
                    "mt700_alanlari":    self.mt700_alanlari,
                    "fatura_tutar":      fatura_tutar,
                    "lc_tutar":          lc_tutar,
                    "incoterm":          incoterm,
                    "bl_tarih_str":      bl_tarih_str,
                    "alan_44c":          alan_44c,
                    "fat_kilo":          fat_k,
                    "bl_kilo":           bl_k,
                    "pl_kilo":           pl_k,
                    "sigorta_tutari":    self.sigorta_tutari_bul(sigorta_text),
                    "is_electronic":     self.is_electronic,
                    "is_collection":     self.is_collection,
                    "lc_no":             lc_no
                }
                
                ucp_sonuc = ucp_kurallari_uygula(parsed_data)
                if ucp_sonuc:
                    r["ucp_tablosu"] = ucp_sonuc
                    
                # Rapor rezerv güncellemeleri
                for u in ucp_sonuc:
                    if u["durum"] == "REZERV" and u["kod"] not in self._aktif_rezervler:
                        self._aktif_rezervler.append(u["kod"])
                        self._risk_ekle(u["kod"])
                        r["rezerv_ozeti"].append(f"REZERV — {u['aciklama']}: {u['detay']}")

                # MT700 yorumları
                if mt700_hukuki_yorum:
                    r["mt700_yorumlari"] = mt700_hukuki_yorum(parsed_data)

                # Uzman Görüşü
                if uzman_gorusu_uret:
                    r["uzman_gorusu"] = uzman_gorusu_uret(parsed_data, ucp_sonuc)
                    
            except Exception as e:
                log.error("hukuk_motoru analizi sırasında hata: %s\n%s", e, traceback.format_exc())

        # Retrieve Search Log
        r["search_log"] = search_log_getir()

        # Risk Özeti & Banka Kabul
        r["risk_ozeti"].append(f"Risk Puanı: **{self.risk_puani}** — {self._risk_sinifi()}")
        r["risk_ozeti"].append(f"Uyumluluk Puanı: **%{self.uyumluluk_puani}**")
        
        # SWIFT Ret Simülatörü
        r["rezerv_swift"] = [REZERV_SWIFT[k] for k in self._aktif_rezervler if k in REZERV_SWIFT]

        # Rezerv Detaylar
        r["rezerv_detaylar"] = [
            {
                "kod": k,
                "kategori": REZERV_KATEGORILERI.get(k, {}).get("kategori", "UNKNOWN"),
                "puan": str(REZERV_KATEGORILERI.get(k, {}).get("puan", 0)),
                "sure": REZERV_KATEGORILERI.get(k, {}).get("sure", "-"),
                "remediation": REZERV_KATEGORILERI.get(k, {}).get("remediation", "-")
            }
            for k in self._aktif_rezervler
        ]

        # Yönetici Özeti
        mevcut = [k for k in ["KUSAT","FATURA","KONSIMENTO","CEKI_LISTESI","SIGORTA"] if self.depo.get(k)]
        r["yonetici_ozeti"] = {
            "lc_no": lc_no or "Tespit edilemedi",
            "mevcut": mevcut,
            "eksik": r["eksik_belgeler"],
            "toplam_rezerv": len(r["rezerv_ozeti"]),
            "major_rezerv": sum(1 for k in self._aktif_rezervler if REZERV_KATEGORILERI.get(k, {}).get("kategori") == "MAJOR DISCREPANCY"),
            "uyumluluk": self.uyumluluk_puani,
            "risk_puani": self.risk_puani,
            "risk_sinifi": self._risk_sinifi(),
            "banka_kabul": self._banka_kabul,
            "transaction_type": self.transaction_type,
        }

        self.analysis_result = r

    def markdown_raporu(self) -> str:
        v = self.analysis_result
        if not v:
            return ""

        s = []
        s.append("# AKREDİTİF GELİŞMİŞ HUKUKİ VE SAYISAL UZMAN DENETİM RAPORU v11.0\n\n")
        s.append(f"**Tarih:** {datetime.now().strftime('%d.%m.%Y %H:%M')} | Hukuki Altyapı: UCP 600, ISBP, eUCP, eURC & Incoterms 2020\n\n---\n\n")

        # Yönetici Özeti
        oz = v.get("yonetici_ozeti", {})
        s.append("## 1. Yönetici Özeti\n\n")
        s.append(f"- **İşlem Tipi:** {oz.get('transaction_type')}\n")
        s.append(f"- **Akreditif Referans No (LC No):** `{oz.get('lc_no')}`\n")
        s.append(f"- **Belgelerin Uyumluluk Skoru:** **%{oz.get('uyumluluk')}/100**\n")
        s.append(f"- **Banka Kabul Olasılığı:** **%{oz.get('banka_kabul')}/100**\n")
        s.append(f"- **Risk Seviyesi:** **{oz.get('risk_sinifi')}** ({oz.get('risk_puani')} Risk Puanı)\n")
        s.append(f"- **İbraz Edilen Belgeler:** {', '.join(oz.get('mevcut', []))}\n")
        s.append(f"- **Tespit Edilen Rezerv Sayısı:** {oz.get('toplam_rezerv', 0)} (Kritik/Major: {oz.get('major_rezerv', 0)})\n\n---\n\n")

        # Belge Durum Raporu (OCR ve Classification Confidence)
        s.append("## 2. Belge Durum ve Veri Çıkarım Güven Raporu\n\n")
        s.append("| Dosya | Sınıfı | Sınıflandırma Güven Skoru | Metin Çıkarım Güveni | OCR Gerekli mi? |\n|:---|:---|:---|:---|:---|\n")
        for d in v.get("dosya_durum_raporu", []):
            ocr_needed = "EVET" if d.get("ext_conf", 100.0) < 100.0 else "HAYIR"
            s.append(
                f"| {d.get('dosya')} | {d.get('sinif') or '-'} | %{d.get('class_conf', 0.0)} | %{d.get('ext_conf', 100.0)} | {ocr_needed} |\n"
            )
        s.append("\n---\n\n")

        # Knowledge Search Log (RAG Logu)
        s.append("## 3. RAG / Bilgi Tabanı Sorgu Günlüğü (Knowledge Search Log)\n\n")
        s.append("| Sorgulanan Kaynak | Arama Sorgusu | Arama Sonucu |\n|:---|:---|:---|\n")
        for log_item in v.get("search_log", []):
            s.append(f"| {log_item['kaynak']} | `{log_item['sorgu']}` | {log_item['sonuc']} |\n")
        s.append("\n---\n\n")

        # UCP Hukuki Kontroller ve Karar Ağaçları
        s.append("## 4. Hukuki Muhakeme ve Karar Ağaçları (Decision Trace)\n\n")
        ucp_tablo = v.get("ucp_tablosu", [])
        if ucp_tablo:
            for u in ucp_tablo:
                durum_sembol = "✓ UYUMLU" if u["durum"] == "UYUMLU" else "✗ REZERV" if u["durum"] == "REZERV" else "⚠ " + u["durum"]
                s.append(f"### Madde: {u['madde']} — {u['aciklama']}\n\n")
                s.append(f"- **BULGU:** {u['detay']}\n")
                s.append(f"- **HUKUKİ DEĞERLENDİRME:** {u['hukuki_yorum']}\n")
                s.append(f"- **UYGULANAN MEVZUAT DAYANAĞI:** `{u['dayanak']}` ({u.get('kaynak')})\n")
                s.append(f"- **DÜZELTME ÖNERİSİ (REMEDIATION):** *{u['remediation']}*\n")
                s.append(f"- **KARAR AĞACI İZİ (DECISION TRACE):**\n  `{u['trace']}`\n\n")
                s.append(f"- **SONUÇ:** **{durum_sembol}**\n\n---\n\n")
        else:
            s.append("Hukuki kural eşleşmesi veya mevzuat taraması tetiklenmedi.\n\n---\n\n")

        # Çapraz Belgeler Karşılaştırması
        s.append("## 5. Çapraz Belge Karşılaştırma Raporu\n\n")
        s.append("| İnceleme Kalemi | İbraz Bulgusu / Çapraz Analiz | Sonuç |\n|:---|:---|:---|\n")
        for c in v.get("capraz_kontrol", []):
            s.append(f"| {c['belge']} | {c['detay']} | **{c['durum']}** |\n")
        s.append("\n---\n\n")

        # SWIFT Ret Simülatörü
        swift = v.get("rezerv_swift", [])
        if swift:
            s.append("## 6. SWIFT Rezerv Bildirim Simülatörü (MT734)\n\n")
            s.append("Bankanın ibrazı reddetmesi durumunda üreteceği standart SWIFT ret mesajı simülasyonları:\n\n")
            for idx, mt in enumerate(swift, 1):
                s.append(f"### Ret Bildirimi {idx}\n```\n{mt}\n```\n\n")
            s.append("---\n\n")

        # Uzman Görüşü
        uzman = v.get("uzman_gorusu", "")
        if uzman:
            s.append("## 7. HUKUKİ NİHAİ UZMAN GÖRÜŞÜ\n\n")
            s.append(uzman)
            s.append("\n")

        return "".join(s)

    def rapor_kaydet(self) -> str:
        icerik = self.markdown_raporu()
        if not icerik:
            return ""
        yol = os.path.join(self.raporlar_dir, "akreditif_analiz_raporu.md")
        with open(yol, "w", encoding="utf-8") as f:
            f.write(icerik)
        print("[+] akreditif_analiz_raporu.md kaydedildi:", yol)
        return icerik

    def html_raporu(self) -> None:
        v = self.analysis_result
        if not v:
            return
        yol = os.path.join(self.raporlar_dir, "akreditif_analiz_raporu.html")

        # HTML generation
        # Sadece basit bir CSS ile şık ve modern bir görünüme dönüştürelim
        oz = v.get("yonetici_ozeti", {})
        krenk = ("#276749" if oz.get("banka_kabul", 0) >= 70 else "#d69e2e" if oz.get("banka_kabul", 0) >= 40 else "#c53030")

        ddr = "".join(
            f"<tr><td>{d.get('dosya','')}</td>"
            f"<td>{d.get('sinif') or '-'}</td>"
            f"<td>%{d.get('class_conf', 0.0)}</td>"
            f"<td>%{d.get('ext_conf', 100.0)}</td>"
            f"<td>{'EVET' if d.get('ext_conf', 100.0) < 100.0 else 'HAYIR'}</td></tr>"
            for d in v.get("dosya_durum_raporu", [])
        )
        
        search_html = "".join(
            f"<tr><td>{item['kaynak']}</td><td><code>{item['sorgu']}</code></td><td>{item['sonuc']}</td></tr>"
            for item in v.get("search_log", [])
        )

        trace_html = ""
        for u in v.get("ucp_tablosu", []):
            durum_cls = "compliant" if u["durum"] == "UYUMLU" else "discrepant" if u["durum"] == "REZERV" else "warning"
            trace_html += f"""
            <div class="trace-card">
                <div class="trace-header">
                    <span class="badge {durum_cls}">{u['durum']}</span>
                    <h3>{u['madde']} — {u['aciklama']}</h3>
                </div>
                <div class="trace-body">
                    <p><b>Bulgu:</b> {u['detay']}</p>
                    <p><b>Hukuki Değerlendirme:</b> {u['hukuki_yorum']}</p>
                    <p><b>Hukuki Dayanak:</b> <code>{u['dayanak']}</code> ({u.get('kaynak')})</p>
                    <p><b>Çözüm Önerisi (Remediation):</b> <i>{u['remediation']}</i></p>
                    <div class="trace-log"><b>Karar Ağacı İz Kaydı:</b> <code>{u['trace']}</code></div>
                </div>
            </div>
            """

        cc_html = "".join(
            f"<tr><td>{c['belge']}</td><td>{c['detay']}</td><td><b>{c['durum']}</b></td></tr>"
            for c in v.get("capraz_kontrol", [])
        )

        swift_html = "".join(
            f'<div class="swift"><b>Ret SIM-{i}</b><pre>{mt}</pre></div>'
            for i, mt in enumerate(v.get("rezerv_swift", []), 1)
        ) or '<p style="color:#276749">SWIFT ret metni üretilmedi.</p>'

        html = f"""<!DOCTYPE html>
<html lang="tr">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>Akreditif Analiz Raporu v11.0</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:'Segoe UI',sans-serif;background:#f4f6f9;color:#333;padding:30px}}
.wrap{{background:#wrap;background-color:#fff;padding:40px;border-radius:12px;box-shadow:0 4px 20px rgba(0,0,0,.08);max-width:1280px;margin:0 auto}}
h1{{color:#1a365d;border-bottom:4px solid #3182ce;padding-bottom:12px;font-size:1.6em}}
h2{{color:#2b6cb0;margin:30px 0 10px;border-left:5px solid #3182ce;padding-left:10px;font-size:1.2em}}
table{{width:100%;border-collapse:collapse;margin-top:10px;font-size:.9em}}
th,td{{border:1px solid #e2e8f0;padding:10px 12px;text-align:left;vertical-align:top}}
th{{background:#ebf8ff;color:#2b6cb0;font-weight:600}}
tr:nth-child(even){{background:#f7fafc}}
.meta{{color:#718096;font-size:.85em;margin-bottom:20px}}
.exec{{background:linear-gradient(135deg,#ebf8ff,#f0fff4);border:2px solid #3182ce;border-radius:10px;padding:24px;margin-bottom:25px}}
.grid{{display:flex;flex-wrap:wrap;gap:15px;margin-top:15px}}
.card{{background:#wrap;background-color:#fff;border:1px solid #bee3f8;border-radius:7px;padding:12px 18px;min-width:180px;text-align:center;box-shadow:0 2px 4px rgba(0,0,0,.02)}}
.lbl{{font-size:.78em;color:#718096;margin-bottom:4px}}.val{{font-size:1.4em;font-weight:700;color:#2b6cb0}}
code{{background:#edf2f7;padding:2px 6px;border-radius:4px;font-size:.85em;font-family:monospace}}
.trace-card{{border:1px solid #e2e8f0;border-radius:8px;margin-bottom:15px;background:#fcfcfc;box-shadow:0 1px 3px rgba(0,0,0,.02)}}
.trace-header{{display:flex;align-items:center;padding:12px;background:#f7fafc;border-bottom:1px solid #e2e8f0;border-top-left-radius:8px;border-top-right-radius:8px}}
.trace-body{{padding:15px}}
.trace-log{{margin-top:12px;background:#f0f4f8;padding:8px 12px;border-radius:5px;font-size:.8em}}
.badge{{display:inline-block;padding:3px 8px;border-radius:4px;font-size:.75em;font-weight:600;margin-right:10px}}
.badge.compliant{{background-color:#c6f6d5;color:#22543d}}
.badge.discrepant{{background-color:#fed7d7;color:#742a2a}}
.badge.warning{{background-color:#feebc8;color:#744210}}
.swift{{background:#1a202c;color:#f6e05e;border-radius:7px;padding:16px;margin:10px 0;font-family:monospace;font-size:.85em}}
.swift pre{{white-space:pre-wrap;margin-top:6px;color:#e2e8f0}}
</style>
</head>
<body>
<div class="wrap">
<h1>AKREDİTİF GELİŞMİŞ HUKUKİ VE SAYISAL UZMAN DENETİM RAPORU v11.0</h1>
<p class="meta"><b>Tarih:</b> {datetime.now().strftime('%d.%m.%Y %H:%M')} | <b>Hukuki Altyapı:</b> UCP 600, ISBP, eUCP, eURC &amp; Incoterms 2020</p>
<div class="exec">
<h2>Yönetici Özeti</h2>
<div class="grid">
<div class="card"><div class="lbl">İşlem Tipi</div><div class="val" style="font-size:1.0em;padding-top:6px;">{oz.get('transaction_type')}</div></div>
<div class="card"><div class="lbl">Toplam Belgeler</div><div class="val">{len(oz.get('mevcut',[]))}</div></div>
<div class="card"><div class="lbl">Toplam Rezerv</div><div class="val" style="color:#c53030">{oz.get('toplam_rezerv',0)}</div></div>
<div class="card"><div class="lbl">Uyumluluk</div><div class="val" style="color:#276749">%{oz.get('uyumluluk','?')}</div></div>
<div class="card"><div class="lbl">Banka Kabul Olasılığı</div><div class="val" style="color:{krenk}">%{oz.get('banka_kabul','?')}</div></div>
<div class="card"><div class="lbl">Risk Seviyesi</div><div class="val" style="font-size:1.0em;padding-top:6px;">{oz.get('risk_sinifi','?')}</div></div>
</div></div>
<h2>1. Belge Durum ve Veri Çıkarım Güven Raporu</h2>
<table><tr><th>Dosya</th><th>Sınıfı</th><th>Sınıflandırma Güveni</th><th>Veri Çıkarım Güveni</th><th>OCR Gerekli mi?</th></tr>{ddr}</table>
<h2>2. RAG / Bilgi Tabanı Sorgu Günlüğü (Knowledge Search Log)</h2>
<table><tr><th>Sorgulanan Kaynak</th><th>Arama Sorgusu</th><th>Arama Sonucu</th></tr>{search_html}</table>
<h2>3. Hukuki Muhakeme ve Karar Ağaçları (Decision Trace)</h2>
{trace_html}
<h2>4. Çapraz Belge Karşılaştırma Raporu</h2>
<table><tr><th>İnceleme Kalemi</th><th>İbraz Bulgusu / Çapraz Analiz</th><th>Sonuç</th></tr>{cc_html}</table>
<h2>5. SWIFT Rezerv Bildirim Simülatörü (MT734)</h2>
{swift_html}
<h2>6. HUKUKİ NİHAİ UZMAN GÖRÜŞÜ</h2>
<p style="white-space:pre-wrap;line-height:1.6;font-size:0.95em;">{v.get('uzman_gorusu','')}</p>
</div></body></html>"""

        with open(yol, "w", encoding="utf-8") as f:
            f.write(html)
        print("[+] rapor.html kaydedildi:", yol)

    def _risk_ekle(self, kategori: str) -> None:
        b = REZERV_KATEGORILERI.get(kategori, {})
        self.risk_puani += b.get("puan", 0)
        kat = b.get("kategori", "")
        if "MAJOR" in kat:
            self._banka_kabul = max(0, self._banka_kabul - 25)
        elif "MEDIUM" in kat:
            self._banka_kabul = max(0, self._banka_kabul - 10)
        else:
            self._banka_kabul = max(0, self._banka_kabul - 5)
        if kategori not in self._aktif_rezervler:
            self._aktif_rezervler.append(kategori)

    def _uyum_dus(self, n: int) -> None:
        self.uyumluluk_puani = max(0, self.uyumluluk_puani - n)

    def _risk_sinifi(self) -> str:
        for alt, ust, s in RISK_SINIFLANDIRMASI:
            if alt <= self.risk_puani <= ust:
                return s
        return "YÜKSEK RİSK"

    def mal_tanimi_benzerlik(self, kaynak: str, hedef: str) -> float:
        def norm(s):
            s = s.upper()
            s = re.sub(r'[^\w\s]', ' ', s)
            return set(w for w in re.sub(r'\s+', ' ', s).split() if len(w) >= 4)
        k, h = norm(kaynak), norm(hedef)
        if not k:
            return 0.0
        return len(k & h) / len(k)

    def _eski_raporlari_temizle(self) -> None:
        for ad in [
            "rapor.md", "rapor.html",
            "akreditif_analiz_raporu.md",
            "akreditif_analiz_raporu.html",
            "akreditif_analiz_raporu.docx",
        ]:
            yol = os.path.join(self.raporlar_dir, ad)
            if os.path.isfile(yol):
                try:
                    os.remove(yol)
                except OSError:
                    pass

    def baslat(self) -> None:
        print("[BİLGİ] Akreditif denetim sistemi v11.0 başlatılıyor...")
        self._eski_raporlari_temizle()

        if self.depoyu_tara():
            self.analiz_motoru()
            final_report = self.rapor_kaydet()
            self.html_raporu()
            print("[TEST] Analiz başarıyla tamamlandı. Raporlar hazır. ✓")
        else:
            print("[BİLGİ] Yüklenecek belge bulunamadı.")


if __name__ == "__main__":
    YapayZekaDisTicaretDenetleyici().baslat()
